import streamlit as st
from docx import Document
import random
import os
import mammoth

# --- Replaces placeholders across styled runs, preserving format ---
def replace_text_format_preserved(doc, replacements):
    def replace_in_runs(runs, replacements):
        full_text = ''.join(run.text for run in runs)
        for key, value in replacements.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))
                for run in runs:
                    run.text = ''
                if runs:
                    runs[0].text = full_text
                break

    for para in doc.paragraphs:
        replace_in_runs(para.runs, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para.runs, replacements)

# --- Generates the output COA .docx ---
def generate_docx(data, template_path="template.docx", output_path="generated_coa.docx"):
    doc = Document(template_path)
    replace_text_format_preserved(doc, data)
    doc.save(output_path)
    return output_path

# --- Optional HTML preview using mammoth ---
def docx_to_html(docx_path):
    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        return result.value

# --- Component calculations based on moisture ---
def calculate_components(moisture):
    remaining = 100 - moisture
    gum = round(random.uniform(81, min(85, remaining - 1.5)), 2)
    remaining -= gum
    protein = round(min(5, remaining * 0.2), 2)
    remaining -= protein
    ash = round(min(1, remaining * 0.2), 2)
    remaining -= ash
    air = round(min(6, remaining * 0.5), 2)
    remaining -= air
    fat = round(remaining, 2)
    return gum, protein, ash, air, fat

# --- UI Starts ---
st.set_page_config(page_title="COA Generator", layout="wide")
st.title("🧪 COA Document Generator (Code-Based Template)")

with st.form("coa_form"):
    code = st.selectbox(
        "Select Product Code Range",
        [f"{i}-{i+500}" for i in range(500, 10001, 500)]
    )
    date = st.text_input("Date (e.g., JULY 2025)")
    batch_no = st.text_input("Batch Number")
    best_before = st.text_input("Best Before (e.g., JULY 2027)")
    moisture = st.number_input("Moisture (%)", min_value=0.0, max_value=100.0, step=0.01, value=10.0)
    ph = st.text_input("pH Level (e.g., 6.7)")
    mesh_200 = st.text_input("200 Mesh (%)")
    viscosity_2h = st.text_input("Viscosity After 2 Hours (CPS)")
    viscosity_24h = st.text_input("Viscosity After 24 Hours (CPS)")
    submitted = st.form_submit_button("Generate COA")

if submitted:
    try:
        # Ensure template exists
        template_path = f"COA {code}.docx"
        if not os.path.exists(template_path):
            st.error(f"Template file 'COA {code}.docx' not found!")
        else:
            gum, protein, ash, air, fat = calculate_components(moisture)

            data = {
                "DATE": date,
                "BATCH_NO": batch_no,
                "BEST_BEFORE": best_before,
                "MOISTURE": f"{moisture}%",
                "PH": ph,
                "MESH_200": f"{mesh_200}%",
                "VISCOSITY_2H": viscosity_2h,
                "VISCOSITY_24H": viscosity_24h,
                "GUM_CONTENT": f"{gum}%",
                "PROTEIN": f"{protein}%",
                "ASH_CONTENT": f"{ash}%",
                "AIR": f"{air}%",
                "FAT": f"{fat}%"
            }

            output_path = "generated_coa.docx"
            generate_docx(data, template_path=template_path, output_path=output_path)

            # Preview
            try:
                html = docx_to_html(output_path)
                st.subheader("📄 Preview")
                st.components.v1.html(f"<div style='padding:15px'>{html}</div>", height=700, scrolling=True)
            except:
                st.warning("Preview failed. You can still download the file below.")

            with open(output_path, "rb") as file:
                st.download_button("📥 Download COA (DOCX)", file, file_name="COA_Generated.docx")
    except Exception as e:
        st.error(f"❌ Error: {str(e)}")
