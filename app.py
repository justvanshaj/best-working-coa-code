import streamlit as st
from docx import Document
from docx.shared import RGBColor
from datetime import datetime
import calendar
import random
import os
import pandas as pd
import io
import mammoth
import zipfile

# --- Style-preserving text replacement ---
def advanced_replace_text_preserving_style(doc, replacements):
    def replace_in_paragraph(paragraph):
        runs = paragraph.runs
        full_text = ''.join(run.text for run in runs)
        for key, value in replacements.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in full_text:
                new_runs = []
                accumulated = ""
                for run in runs:
                    accumulated += run.text
                    new_runs.append(run)
                    if placeholder in accumulated:
                        style_run = next((r for r in new_runs if placeholder in r.text), new_runs[0])
                        font = style_run.font
                        accumulated = accumulated.replace(placeholder, value)
                        for r in new_runs:
                            r.text = ''
                        if new_runs:
                            new_run = new_runs[0]
                            new_run.text = accumulated
                            new_run.font.name = font.name
                            new_run.font.size = font.size
                            new_run.font.bold = font.bold
                            new_run.font.italic = font.italic
                            new_run.font.underline = font.underline
                            new_run.font.color.rgb = font.color.rgb
                        break

    for para in doc.paragraphs:
        replace_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

# --- Generate DOCX file ---
def generate_docx(data, template_path, output_path):
    doc = Document(template_path)
    advanced_replace_text_preserving_style(doc, data)
    doc.save(output_path)

# --- Convert DOCX to HTML for preview ---
def docx_to_html(path):
    with open(path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        return result.value

# --- Updated: Calculate components from moisture ---
def calculate_components(moisture):
    total = 100
    remaining = total - moisture

    gum = round(random.uniform(81, min(88, remaining - 1.5)), 2)
    remaining -= gum

    # Initial base values
    base_protein = min(4, remaining * 0.2)
    base_ash = min(0.7, remaining * 0.2)
    base_air = min(3.5, remaining * 0.5)
    base_fat = min(0.8, remaining)

    base_total = base_protein + base_ash + base_air + base_fat

    if round(base_total, 2) <= round(remaining, 2):
        protein = round(base_protein, 2)
        ash = round(base_ash, 2)
        air = round(base_air, 2)
        fat = round(remaining - (protein + ash + air), 2)
        fat = min(fat, 0.8)

        leftover = round(remaining - (protein + ash + air + fat), 2)
        if leftover > 0 and (protein + ash + air) > 0:
            scale = remaining / (protein + ash + air)
            protein = round(protein * scale, 2)
            ash = round(ash * scale, 2)
            air = round(air * scale, 2)
            fat = round(remaining - (protein + ash + air), 2)
    else:
        scale = remaining / base_total
        protein = round(base_protein * scale, 2)
        ash = round(base_ash * scale, 2)
        air = round(base_air * scale, 2)
        fat = round(base_fat * scale, 2)
        fat = min(fat, 0.8)

        subtotal = protein + ash + air + fat
        if subtotal < remaining and (protein + ash + air) > 0:
            extra = remaining - subtotal
            protein += round(extra * (protein / (protein + ash + air)), 2)
            ash += round(extra * (ash / (protein + ash + air)), 2)
            air += round(extra * (air / (protein + ash + air)), 2)
            fat = round(remaining - (protein + ash + air), 2)

    return gum, protein, ash, air, fat

# --- Streamlit UI ---
st.set_page_config("Bulk COA Generator", layout="wide")
st.title("📥 Bulk COA Generator from Excel")

uploaded_file = st.file_uploader("Upload Excel File", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)
    st.success(f"✅ Uploaded {len(df)} row(s)")

    coa_files = []
    temp_dir = "generated_coas"
    os.makedirs(temp_dir, exist_ok=True)

    for idx, row in df.iterrows():
        try:
            code = str(row["Code"]).strip()
            date = str(row["Date"]).strip()
            batch = str(row["Batch No"]).strip()
            moisture = float(row["Moisture"])
            ph = str(row["pH"]).strip()
            mesh = str(row["200 Mesh"]).strip()
            vis2h = str(row["Viscosity 2H"]).strip()
            vis24h = str(row["Viscosity 24H"]).strip()

            # Calculate Best Before
            try:
                dt = datetime.strptime(date.strip(), "%B %Y")
                year = dt.year + 2
                month = dt.month - 1
                if month == 0:
                    month = 12
                    year -= 1
                best_before = f"{calendar.month_name[month].upper()} {year}"
            except:
                best_before = "N/A"

            gum, protein, ash, air, fat = calculate_components(moisture)

            data = {
                "DATE": date,
                "BATCH_NO": batch,
                "BEST_BEFORE": best_before,
                "MOISTURE": f"{moisture}%",
                "PH": ph,
                "MESH_200": f"{mesh}%",
                "VISCOSITY_2H": vis2h,
                "VISCOSITY_24H": vis24h,
                "GUM_CONTENT": f"{gum}%",
                "PROTEIN": f"{protein}%",
                "ASH_CONTENT": f"{ash}%",
                "AIR": f"{air}%",
                "FAT": f"{fat}%"
            }

            safe_batch = batch.replace("/", "_").replace("\\", "_").replace(" ", "_")
            filename = f"COA-{safe_batch}-{code}.docx"
            template = f"COA {code}.docx"
            output_path = os.path.join(temp_dir, filename)

            if os.path.exists(template):
                generate_docx(data, template, output_path)
                coa_files.append(output_path)
                st.success(f"✅ Generated: {filename}")
            else:
                st.warning(f"⚠️ Missing template: COA {code}.docx (skipped row {idx+1})")

        except Exception as e:
            st.error(f"❌ Error processing row {idx+1}: {str(e)}")

    if coa_files:
        st.subheader("📄 Download Generated COAs")

        for path in coa_files:
            with open(path, "rb") as f:
                st.download_button(
                    label=f"⬇️ Download {os.path.basename(path)}",
                    data=f,
                    file_name=os.path.basename(path),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        # ZIP all files
        zip_path = os.path.join(temp_dir, "All_COAs.zip")
        with zipfile.ZipFile(zip_path, "w") as zipf:
            for f in coa_files:
                zipf.write(f, arcname=os.path.basename(f))

        with open(zip_path, "rb") as z:
            st.download_button("📦 Download All as ZIP", z, file_name="All_COAs.zip")
